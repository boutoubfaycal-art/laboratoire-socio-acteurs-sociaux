import docx
import os

def extract_full_text(doc):
    """Extraire tout le texte d'un document Word"""
    full_text = []
    
    # Extraire les paragraphes
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extraire les tableaux
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(" | ".join(row_text))
    
    return full_text

# Lire toutes les parties
print("=== EXTRACTION COMPLÈTE DU CONTENU ===\n")

parts_data = {}

for i in range(1, 8):
    # Trouver le fichier correspondant
    matching_files = [f for f in os.listdir('.') if f.startswith(f'Partie {i}') and f.endswith('.docx')]
    
    if matching_files:
        filename = matching_files[0]
        try:
            doc = docx.Document(filename)
            full_content = extract_full_text(doc)
            
            parts_data[i] = {
                'filename': filename,
                'content': full_content
            }
            
            print(f"=== PARTIE {i} ===")
            print(f"Fichier: {filename}")
            print(f"Nombre total d'éléments: {len(full_content)}")
            print("Contenu complet:")
            for j, text in enumerate(full_content):
                print(f"  {j+1:2d}. {text}")
            print()
            
        except Exception as e:
            print(f"Erreur lecture {filename}: {e}")
            parts_data[i] = {'filename': filename, 'content': []}
    else:
        print(f"Aucun fichier trouvé pour la Partie {i}")
        parts_data[i] = {'filename': '', 'content': []}

# Créer le contenu structuré pour les 4 sections
structured_content = {
    "section1_nomination": {
        "title": "تسمية المخبر",
        "title_fr": "Nomination du laboratoire",
        "content": parts_data.get(1, {}).get("content", [])
    },
    "section2_presentation": {
        "title": "تقديم المخبر",
        "title_fr": "Présentation du laboratoire", 
        "content": parts_data.get(2, {}).get("content", [])
    },
    "section3_definition": {
        "title": "تعريف المخبر",
        "title_fr": "Définition du laboratoire",
        "content": parts_data.get(3, {}).get("content", [])
    },
    "section4_equipes_recherche": {
        "title": "فرق البحث",
        "title_fr": "Équipes de recherche",
        "equipes": [
            {
                "numero": 1,
                "titre": parts_data.get(4, {}).get("content", ["الشباب والقيم في المجتمع الجزائري: العمل، السياسة، والممارسات الثقافية"])[0] if parts_data.get(4) and len(parts_data.get(4).get("content", [])) > 0 else "الشباب والقيم في المجتمع الجزائري: العمل، السياسة، والممارسات الثقافية",
                "titre_fr": "Jeunesse et valeurs dans la société algérienne: Travail, politique et pratiques culturelles",
                "contenu": parts_data.get(4, {}).get("content", ["Contenu non disponible - Partie 4 manquante"])
            },
            {
                "numero": 2,
                "titre": parts_data.get(5, {}).get("content", [""])[1] if len(parts_data.get(5, {}).get("content", [])) > 1 else "الرابطة الاجتماعية والأسرة",
                "titre_fr": "Lien social et famille",
                "contenu": parts_data.get(5, {}).get("content", [])
            },
            {
                "numero": 3,
                "titre": parts_data.get(6, {}).get("content", [""])[1] if len(parts_data.get(6, {}).get("content", [])) > 1 else "الفاعلون الاجتماعيون في الوسط الحضري",
                "titre_fr": "Acteurs sociaux en milieu urbain",
                "contenu": parts_data.get(6, {}).get("content", [])
            },
            {
                "numero": 4,
                "titre": parts_data.get(7, {}).get("content", [""])[1] if len(parts_data.get(7, {}).get("content", [])) > 1 else "قيم الشباب الجزائري والبيئة الرقمية",
                "titre_fr": "Valeurs des jeunes algériens et environnement numérique",
                "contenu": parts_data.get(7, {}).get("content", [])
            }
        ]
    }
}

# Sauvegarder le contenu structuré
import json
with open('structured_content.json', 'w', encoding='utf-8') as f:
    json.dump(structured_content, f, ensure_ascii=False, indent=2)

print("=== CONTENU STRUCTURÉ CRÉÉ ===")
print("Sections créées:")
print("1. تسمية المخبر (Nomination)")
print("2. تقديم المخبر (Présentation)")  
print("3. تعريف المخبر (Définition)")
print("4. فرق البحث (Équipes de recherche) - regroupant les parties 4, 5, 6, 7")
print(f"\nContenu sauvegardé dans 'structured_content.json'")
